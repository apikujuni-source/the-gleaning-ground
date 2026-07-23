from __future__ import annotations

from pathlib import Path
from textwrap import wrap

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from reportlab.lib import colors
from reportlab.lib.pagesizes import portrait
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas


OUTPUT = Path(__file__).parent / "output"
OUTPUT.mkdir(parents=True, exist_ok=True)

DOCX_PATH = OUTPUT / "The-Divine-Blueprint-Companion-Journal-Editable.docx"
PDF_PATH = OUTPUT / "The-Divine-Blueprint-Companion-Journal-Print-Ready.pdf"
NOTES_PATH = OUTPUT / "The-Divine-Blueprint-Companion-Journal-Production-Notes.txt"

NAVY = "0B2036"
NAVY_RGB = RGBColor(11, 32, 54)
GOLD = "B8862D"
GOLD_RGB = RGBColor(184, 134, 45)
CREAM = "F7F1E5"
CREAM_RGB = RGBColor(247, 241, 229)
INK = "16283A"
INK_RGB = RGBColor(22, 40, 58)
MUTED = "6D6A63"
MUTED_RGB = RGBColor(109, 106, 99)
WHITE_RGB = RGBColor(255, 255, 255)

PAGE_W = 7 * inch
PAGE_H = 10 * inch

CHAPTERS = [
    {
        "num": 1,
        "title": "The Light of Men",
        "display": "THE LIGHT\nof Men",
        "key_ref": "John 1:4",
        "key_text": "In him was life; and the life was the light of men.",
        "supporting": ["Genesis 1:1–3", "John 1:1–5", "John 1:9–13", "2 Corinthians 5:17"],
        "objectives": [
            "Understand why God's first act in creation is to bring light.",
            "Identify areas of darkness, emptiness, and disorder in your own life.",
            "Recognize how the Holy Spirit prepares the heart for transformation.",
            "Begin cooperating with God's recreative work through Christ.",
        ],
        "prayer": "Father, I surrender my assumptions, fears, and pride. Let Your Word expose every area of darkness within me. As I read, let me encounter Jesus, the Light of men. Transform me from the inside out. Amen.",
        "focus": ["Darkness", "Light", "Identity", "Recreation"],
        "pathway": ["Darkness", "Light", "Identity", "Purpose", "Fruitfulness"],
        "inventory": [
            "I can name the areas of my life that still feel disordered.",
            "I regularly invite the Holy Spirit to search my heart.",
            "I receive Jesus as the true source of light and life.",
            "My identity is increasingly rooted in what God says about me.",
            "I am taking practical steps toward a fruitful life.",
        ],
        "areas": [
            ("WITHOUT FORM", "God's light exposes disorder and begins to establish divine order.", ["What areas of my life feel disordered?", "What identities have I accepted that did not come from God?", "What would God's order look like here?"]),
            ("VOID", "Emptiness cannot be permanently filled by achievement, possessions, or recognition.", ["Where do I feel spiritually empty?", "What have I tried to use to fill that emptiness?", "What invitation from Christ am I ready to receive?"]),
            ("DARKNESS", "Darkness distorts judgment, but the Light reveals truth and a new beginning.", ["Where have I been unable to see clearly?", "What truth has God recently revealed to me?", "What response of obedience will bring that truth into practice?"]),
        ],
        "meditation_ref": "Genesis 1:3",
        "meditation_text": "And God said, Let there be light: and there was light.",
        "observe": ["What does the verse reveal about God's authority?", "What does light accomplish?", "What darkness does this Word confront in me?", "What phrase should I carry into prayer?"],
        "story": ["Where were you before Christ brought light?", "What changed when you encountered Him?", "What is still changing as you walk in His light?"],
        "practices": ["Read John 1 every morning.", "Pray for fifteen minutes before checking your phone.", "Memorize John 1:4.", "Spend one hour without distractions, listening to God.", "Write one insight from Scripture each day."],
        "declarations": ["I reject darkness.", "I receive the Light of Christ.", "I belong to God.", "I am no longer without form.", "God is giving me His identity.", "My life has purpose.", "His Word recreates me daily.", "I am becoming who God designed me to be."],
        "checkpoint": ["I understand God's purpose for salvation.", "I know my identity in Christ.", "I spend intentional time with God.", "I recognize areas of needed transformation.", "I desire transformation more than information."],
        "milestone": "The greatest thing God showed me in this chapter was…",
        "different": "One thing I will do differently beginning today is…",
    },
    {
        "num": 2,
        "title": "If Children, Then Heirs",
        "display": "IF CHILDREN,\nTHEN HEIRS",
        "key_ref": "Romans 8:16–17",
        "key_text": "The Spirit itself beareth witness with our spirit, that we are the children of God: and if children, then heirs.",
        "supporting": ["Romans 8:14–17", "Galatians 4:4–7", "Colossians 1:12–14", "Ephesians 1:3–6"],
        "objectives": [
            "Understand adoption, redemption, and translation into God's kingdom.",
            "Receive the witness of the Spirit concerning your belonging.",
            "Recognize the inheritance available to God's children in Christ.",
            "Begin living from acceptance rather than striving for acceptance.",
        ],
        "prayer": "Father, open my heart to the wonder of adoption. Deliver me from the mentality of an outsider and teach me to live as Your beloved child. Let the Spirit of adoption silence fear and awaken holy confidence. Amen.",
        "focus": ["Adoption", "Redemption", "Belonging", "Inheritance"],
        "pathway": ["Delivered", "Adopted", "Named", "Heir", "Steward"],
        "inventory": [
            "I relate to God as a loving Father, not merely as a distant authority.",
            "I believe Christ has fully paid the price for my redemption.",
            "I am learning what belongs to me in Christ.",
            "I can receive correction without questioning whether I am loved.",
            "I use my inheritance to serve God's purposes, not merely myself.",
        ],
        "areas": [
            ("DELIVERED", "Redemption breaks the authority of darkness and brings us back to God.", ["What old captivity still shapes my thinking?", "Where do I need to stand in Christ's finished work?", "What evidence of freedom can I thank God for?"]),
            ("ADOPTED", "The Spirit gives inward assurance that we truly belong in the Father's house.", ["When do I feel like an outsider with God?", "What does the Father's welcome mean to me?", "How would secure belonging change my prayers?"]),
            ("HEIRS", "Inheritance is both privilege and responsibility; what God gives must be stewarded.", ["Which spiritual riches have I neglected?", "How can I use what God has given to bless others?", "What responsibility accompanies my inheritance?"]),
        ],
        "meditation_ref": "Galatians 4:7",
        "meditation_text": "Wherefore thou art no more a servant, but a son; and if a son, then an heir of God through Christ.",
        "observe": ["What identity is removed?", "What identity is given?", "Who makes this transition possible?", "How should an heir live today?"],
        "story": ["What has shaped your understanding of fatherhood?", "Where have you experienced God's acceptance?", "What part of your inheritance are you learning to receive?"],
        "practices": ["Address God as Father in your daily prayer.", "List ten gifts of grace you possess in Christ.", "Read Romans 8:1–17 aloud twice this week.", "Reject one recurring thought of spiritual rejection.", "Use one gift or resource to serve another person."],
        "declarations": ["I have been redeemed by Christ.", "I have been brought into God's family.", "The Spirit bears witness that I am God's child.", "Fear does not define my relationship with the Father.", "I am accepted in the Beloved.", "I am an heir of God through Christ.", "My inheritance serves God's purpose.", "I live from belonging, not for belonging."],
        "checkpoint": ["I believe I am fully accepted by God in Christ.", "I understand redemption as a completed work.", "I know several aspects of my inheritance.", "I can identify and resist an orphan mentality.", "I am stewarding what God has placed in my hands."],
        "milestone": "The truth about belonging that most changed my thinking was…",
        "different": "One way I will live as an heir rather than an outsider is…",
    },
    {
        "num": 3,
        "title": "Partakers of His Divine Nature",
        "display": "PARTAKERS OF\nHIS DIVINE NATURE",
        "key_ref": "2 Peter 1:4",
        "key_text": "Whereby are given unto us exceeding great and precious promises: that by these ye might be partakers of the divine nature.",
        "supporting": ["2 Peter 1:3–8", "Galatians 5:22–23", "Romans 12:1–2", "2 Corinthians 3:18"],
        "objectives": [
            "Understand that salvation brings participation in God's life and character.",
            "Recognize the transforming role of God's promises.",
            "Identify character patterns that resist the nature of Christ.",
            "Cultivate spiritual fruit through surrender and disciplined practice.",
        ],
        "prayer": "Lord, do more than improve my behavior—form the life of Christ within me. Through Your promises, renew my desires, motives, speech, and relationships. Make Your character visible in the ordinary places of my life. Amen.",
        "focus": ["Promises", "Nature", "Transformation", "Fruit"],
        "pathway": ["Promise", "Faith", "Practice", "Character", "Fruit"],
        "inventory": [
            "I expect God's Word to transform my character.",
            "I can identify areas where my old nature still dominates.",
            "The fruit of the Spirit is becoming visible in my relationships.",
            "I cooperate with grace through consistent spiritual practices.",
            "My private life increasingly agrees with my public confession.",
        ],
        "areas": [
            ("PROMISES", "God's promises do not merely inform us; they invite us into His life and intention.", ["Which promise is God emphasizing to me?", "What unbelief keeps me from participating in it?", "How can I respond to this promise in faith?"]),
            ("CHARACTER", "Divine nature becomes visible through transformed motives, responses, and habits.", ["Which reaction least resembles Christ?", "What usually triggers that reaction?", "What response would reveal God's nature instead?"]),
            ("FRUIT", "Fruit grows through abiding, pruning, patience, and repeated obedience.", ["Which fruit of the Spirit needs attention?", "What practice would create space for growth?", "Who can help me remain accountable?"]),
        ],
        "meditation_ref": "Galatians 5:22–23",
        "meditation_text": "But the fruit of the Spirit is love, joy, peace, longsuffering, gentleness, goodness, faith, meekness, temperance.",
        "observe": ["Which qualities come from the Spirit?", "Which quality is strongest in me?", "Which quality is weakest?", "What relationship would be changed by this fruit?"],
        "story": ["What character pattern has God already transformed?", "Where is the Spirit currently pruning you?", "What would Christlike maturity look like in your closest relationships?"],
        "practices": ["Meditate on one promise for seven days.", "Choose one fruit of the Spirit as your weekly focus.", "Pause and pray before responding in conflict.", "Ask a trusted believer for honest character feedback.", "Record one evidence of transforming grace each evening."],
        "declarations": ["God has given me precious promises.", "I participate in His life through Christ.", "My old patterns do not have final authority.", "The Spirit is renewing my mind.", "Christ is being formed in me.", "My character can change.", "I bear fruit as I abide in Christ.", "My life will increasingly resemble my Father."],
        "checkpoint": ["I understand participation in the divine nature.", "I regularly meditate on God's promises.", "I can name a current area of character formation.", "I welcome correction and pruning.", "The fruit of the Spirit is becoming more visible."],
        "milestone": "The character quality God is intentionally forming in me is…",
        "different": "One repeated response I will replace with a Christlike response is…",
    },
    {
        "num": 4,
        "title": "But Is Under Tutors",
        "display": "BUT IS UNDER\nTUTORS",
        "key_ref": "Galatians 4:1–2",
        "key_text": "The heir, as long as he is a child, differeth nothing from a servant… but is under tutors and governors until the time appointed of the father.",
        "supporting": ["Galatians 4:1–7", "Hebrews 12:5–11", "Proverbs 12:1", "Luke 16:10–12"],
        "objectives": [
            "Understand the role of tutors, governors, discipline, and spiritual process.",
            "Evaluate your teachability and response to correction.",
            "Recognize the relationship between stewardship and enlargement.",
            "Honor God's timing while faithfully practicing present assignments.",
        ],
        "prayer": "Father, deliver me from impatience and self-importance. Give me a teachable spirit, grace to receive correction, and faithfulness in hidden seasons. Help me honor the people and processes You use to prepare me. Amen.",
        "focus": ["Teachability", "Discipline", "Stewardship", "Timing"],
        "pathway": ["Instruction", "Practice", "Correction", "Trust", "Release"],
        "inventory": [
            "I can receive correction without becoming defensive.",
            "I honor wise spiritual guidance and healthy accountability.",
            "I am faithful with small and unseen responsibilities.",
            "I do not despise seasons of preparation.",
            "I can wait for God's timing without becoming passive.",
        ],
        "areas": [
            ("TUTORS", "God often develops heirs through people who instruct, model, question, and correct.", ["Who has God used to tutor me?", "What lesson have I resisted?", "How can I honor healthy guidance without surrendering discernment?"]),
            ("STEWARDSHIP", "Faithfulness with another person's or a small assignment prepares us for entrusted responsibility.", ["What is currently in my care?", "Where have I been casual or inconsistent?", "What would excellent stewardship require this week?"]),
            ("APPOINTED TIME", "Preparation and release belong to the Father's wisdom, not our impatience.", ["Where am I trying to hurry God's process?", "What can this season produce in me?", "What does faithful waiting look like today?"]),
        ],
        "meditation_ref": "Luke 16:10",
        "meditation_text": "He that is faithful in that which is least is faithful also in much.",
        "observe": ["What does Jesus connect to greater responsibility?", "What is my present 'least' assignment?", "What inconsistency needs attention?", "How can I demonstrate faithfulness today?"],
        "story": ["Describe a season when correction protected or prepared you.", "Who has permission to ask you difficult questions?", "What hidden assignment may be forming your future capacity?"],
        "practices": ["Thank one person who has instructed you.", "Ask for feedback in one area of responsibility.", "Complete one neglected small assignment.", "Study a biblical example of faithful preparation.", "Keep a seven-day stewardship record."],
        "declarations": ["I am teachable before God.", "Correction is not rejection.", "God uses process to prepare me.", "I will be faithful in little.", "I honor wise tutors and governors.", "Hidden seasons are not wasted seasons.", "The Father knows the appointed time.", "I am being prepared for trustworthy service."],
        "checkpoint": ["I welcome biblical correction.", "I have healthy spiritual accountability.", "I am faithful with present responsibilities.", "I understand the value of preparation.", "I trust God's timing for enlargement."],
        "milestone": "The lesson God is teaching me through my present season is…",
        "different": "One small responsibility I will handle with greater faithfulness is…",
    },
    {
        "num": 5,
        "title": "Becoming Sons",
        "display": "BECOMING\nSONS",
        "key_ref": "Romans 8:14",
        "key_text": "For as many as are led by the Spirit of God, they are the sons of God.",
        "supporting": ["1 Corinthians 3:1–3", "Hebrews 5:12–14", "Ephesians 4:13–15", "Romans 8:12–16"],
        "objectives": [
            "Distinguish spiritual birth from spiritual maturity.",
            "Identify patterns of carnality and prolonged infancy.",
            "Embrace responsibility, discernment, and strong spiritual nourishment.",
            "Learn to follow the Spirit as a defining mark of mature sonship.",
        ],
        "prayer": "Father, I do not want to remain spiritually immature. Awaken hunger for truth, courage to accept responsibility, and sensitivity to Your Spirit. Grow me until the character and obedience of Christ are visible in me. Amen.",
        "focus": ["Growth", "Discernment", "Responsibility", "Leading"],
        "pathway": ["New Birth", "Nourishment", "Practice", "Discernment", "Maturity"],
        "inventory": [
            "I can receive strong biblical truth, not only comforting messages.",
            "I take responsibility for my spiritual growth.",
            "I am learning to discern good and evil through practice.",
            "My reactions are becoming less carnal and more Spirit-led.",
            "Other people can depend on me to serve and strengthen them.",
        ],
        "areas": [
            ("NOURISHMENT", "Growth requires moving from milk to strong meat while retaining childlike dependence on God.", ["What truth have I avoided because it felt demanding?", "What spiritual diet currently shapes me?", "What deeper study habit do I need?"]),
            ("RESPONSIBILITY", "Maturity shifts the focus from being continually served to becoming able to serve others.", ["Where am I waiting for others to do what God assigned me?", "Who am I strengthening?", "What responsibility is maturity asking me to accept?"]),
            ("SPIRIT-LED", "Mature sons increasingly recognize, test, and obey the Spirit's leading.", ["How does the Spirit most often guide me?", "What leading have I delayed?", "How can Scripture and community help me discern accurately?"]),
        ],
        "meditation_ref": "Hebrews 5:14",
        "meditation_text": "But strong meat belongeth to them that are of full age… who by reason of use have their senses exercised to discern both good and evil.",
        "observe": ["What belongs to the mature?", "How is discernment developed?", "Where do I need repeated practice?", "What decision requires trained discernment?"],
        "story": ["Where have you clearly grown since your new birth?", "What immature pattern still appears under pressure?", "Who is benefiting from your increasing maturity?"],
        "practices": ["Study one challenging passage without rushing.", "Take responsibility for one ministry or family need.", "Fast from one habit that feeds carnality.", "Record and test impressions you believe are from the Spirit.", "Teach one truth you have practiced."],
        "declarations": ["I will not remain a spiritual infant.", "I hunger for strong truth.", "My senses are being trained through obedience.", "I accept responsibility for growth.", "I put away carnal reactions.", "The Spirit of God leads me.", "I grow into the measure of Christ.", "My maturity will strengthen others."],
        "checkpoint": ["I can identify signs of spiritual immaturity.", "I maintain habits that support growth.", "I am developing biblical discernment.", "I accept responsibility rather than blame others.", "I increasingly follow the Spirit's leading."],
        "milestone": "The clearest evidence of growth God has produced in me is…",
        "different": "One immature pattern I will deliberately replace is…",
    },
    {
        "num": 6,
        "title": "The Cross in the Making of Sons",
        "display": "THE CROSS IN THE\nMAKING OF SONS",
        "key_ref": "Luke 9:23",
        "key_text": "If any man will come after me, let him deny himself, and take up his cross daily, and follow me.",
        "supporting": ["Luke 9:23–24", "Galatians 2:20", "Philippians 2:5–8", "Romans 6:6–13"],
        "objectives": [
            "Understand the cross as God's instrument for dealing with self-rule.",
            "Identify desires, rights, and ambitions that resist Christ's lordship.",
            "Practice daily self-denial, surrender, and obedience.",
            "Follow Jesus through the power of His indwelling life.",
        ],
        "prayer": "Lord Jesus, expose every throne I have built for self. Teach me to deny myself without bitterness, carry my cross without performance, and follow You without negotiation. Let Your life become visible through my surrendered life. Amen.",
        "focus": ["Self-Denial", "Surrender", "Obedience", "Following"],
        "pathway": ["Invitation", "Denial", "Cross", "Following", "Christ's Life"],
        "inventory": [
            "I recognize that discipleship includes costly obedience.",
            "I can release personal rights when love and obedience require it.",
            "I bring ambition and reputation under Christ's lordship.",
            "I obey God even when my emotions resist.",
            "The life of Christ is becoming more visible than self-rule.",
        ],
        "areas": [
            ("DENY SELF", "Self-denial refuses to let personal desire occupy the throne that belongs to Christ.", ["What personal right am I protecting most strongly?", "Where does self insist on control?", "What would surrender look like without resentment?"]),
            ("TAKE UP THE CROSS", "The cross is a daily willingness to let the old nature lose its authority.", ["What recurring situation exposes my flesh?", "What must be put to death in that situation?", "What spiritual practice helps me remain surrendered?"]),
            ("FOLLOW ME", "The goal is not self-improvement but Christ-directed life through the Spirit.", ["Where is Jesus currently leading me?", "What instruction have I negotiated with?", "What next step would put my feet on His path?"]),
        ],
        "meditation_ref": "Galatians 2:20",
        "meditation_text": "I am crucified with Christ: nevertheless I live; yet not I, but Christ liveth in me.",
        "observe": ["What has been crucified?", "Whose life is now expressed?", "What does faith make possible?", "Where do I need to yield the driver's seat?"],
        "story": ["Describe an obedience that cost you something.", "What did God form through that surrender?", "Where is self-rule still competing with Christ's lordship?"],
        "practices": ["Begin each morning with a prayer of surrender.", "Choose hidden service that receives no recognition.", "Fast from one comfort that regularly controls you.", "Obey one clear instruction you have delayed.", "Review your day and name where self or Christ led."],
        "declarations": ["Jesus Christ is Lord of my life.", "I deny the rule of self.", "I take up my cross daily.", "My flesh does not determine my obedience.", "I release the need to control outcomes.", "Christ lives in me.", "The Spirit empowers surrendered living.", "I follow Jesus with my whole life."],
        "checkpoint": ["I understand the daily meaning of the cross.", "I can identify the voice of self-rule.", "I practice surrender in concrete decisions.", "I obey even when obedience is costly.", "Christ's life is increasingly visible in me."],
        "milestone": "The surrender through which God is most deeply forming me is…",
        "different": "One clear instruction I will obey without further negotiation is…",
    },
    {
        "num": 7,
        "title": "Knowledge in the Making of Sons",
        "display": "KNOWLEDGE IN THE\nMAKING OF SONS",
        "key_ref": "Ephesians 1:17–18",
        "key_text": "That the God of our Lord Jesus Christ… may give unto you the spirit of wisdom and revelation in the knowledge of him: the eyes of your understanding being enlightened.",
        "supporting": ["Ephesians 1:15–19", "Colossians 1:9–10", "2 Peter 3:18", "John 17:3"],
        "objectives": [
            "Distinguish information about God from revelational knowledge of God.",
            "Understand the importance of knowing your inheritance in Christ.",
            "Develop habits of Scripture study, meditation, and obedient application.",
            "Pursue intimate knowledge of the Father, Son, and Holy Spirit.",
        ],
        "prayer": "God of our Lord Jesus Christ, give me the Spirit of wisdom and revelation in the knowledge of You. Flood the eyes of my heart with light. Let truth move from my notes into my life, and let knowing You become my highest pursuit. Amen.",
        "focus": ["Revelation", "Inheritance", "The Word", "Intimacy"],
        "pathway": ["Hearing", "Understanding", "Revelation", "Obedience", "Knowledge"],
        "inventory": [
            "My study of Scripture leads to worship and obedience.",
            "I know central truths about my inheritance in Christ.",
            "I ask the Holy Spirit for understanding when I read.",
            "I can distinguish borrowed opinions from personal conviction.",
            "My knowledge of God is becoming relational and experiential.",
        ],
        "areas": [
            ("KNOW YOUR INHERITANCE", "What God has freely given must be understood before it can be faithfully enjoyed and stewarded.", ["Which aspect of my inheritance is unclear?", "What loss results from ignorance in this area?", "What passages should I study carefully?"]),
            ("KNOW HIS PERSON", "Maturity grows through intimate, revelational knowledge of the Father, Son, and Spirit.", ["Which attribute of God is becoming real to me?", "Where has my view of God been distorted?", "What encounter or truth is correcting it?"]),
            ("KNOW HIS WORD", "Truth becomes formative when it is understood, believed, practiced, and retained.", ["Do I read for completion or transformation?", "What truth am I currently practicing?", "How can I move from notes to obedience?"]),
        ],
        "meditation_ref": "John 17:3",
        "meditation_text": "And this is life eternal, that they might know thee the only true God, and Jesus Christ, whom thou hast sent.",
        "observe": ["How does Jesus define eternal life?", "What kind of knowing is implied?", "What competes with knowing God in my priorities?", "What invitation is present in this verse?"],
        "story": ["Describe a truth that moved from information to revelation.", "How did that revelation change your behavior?", "What question about God are you currently carrying?"],
        "practices": ["Study one passage using observation, interpretation, and application.", "Memorize Ephesians 1:17–18.", "Pray Paul's prayers for yourself for seven days.", "Write what Scripture reveals about God's Person.", "Teach one truth only after identifying how you will practice it."],
        "declarations": ["I hunger to know God.", "The Spirit gives wisdom and revelation.", "The eyes of my heart are being enlightened.", "I will understand my inheritance.", "God's Word renews my mind.", "Truth moves me into obedience.", "I grow in grace and knowledge.", "Knowing God is my highest pursuit."],
        "checkpoint": ["I distinguish information from revelation.", "I study Scripture with clear methods.", "I understand key truths about my inheritance.", "I practice what I learn.", "My knowledge of God is increasingly personal."],
        "milestone": "The revelation of God that is reshaping my life is…",
        "different": "One change I will make to my Scripture-learning rhythm is…",
    },
    {
        "num": 8,
        "title": "The Fellowship of the Spirit in the Making of Sons",
        "display": "THE FELLOWSHIP\nOF THE SPIRIT",
        "key_ref": "2 Corinthians 13:14",
        "key_text": "The grace of the Lord Jesus Christ, and the love of God, and the communion of the Holy Ghost, be with you all.",
        "supporting": ["John 16:13", "Acts 8:26–39", "Ephesians 6:18", "Jude 20–21"],
        "objectives": [
            "Understand fellowship with God through the ministry of the Holy Spirit.",
            "Recognize the Spirit's guidance in Scripture, prayer, and daily life.",
            "Develop rhythms of listening, praying, and responding to the Spirit.",
            "Value fellowship with other believers as part of spiritual formation.",
        ],
        "prayer": "Holy Spirit, I welcome Your fellowship. Teach me Christ, lead me into truth, strengthen my prayer, and make my heart responsive to Your promptings. Deliver me from hurried religion and draw me into living communion with the Godhead. Amen.",
        "focus": ["Communion", "Guidance", "Prayer", "Community"],
        "pathway": ["Welcome", "Listen", "Discern", "Respond", "Fellowship"],
        "inventory": [
            "I intentionally welcome the Holy Spirit into my daily rhythms.",
            "I expect the Spirit to illuminate Scripture.",
            "I depend on His help in prayer.",
            "I test impressions through Scripture and mature community.",
            "My fellowship with God produces joy, boldness, and transformation.",
        ],
        "areas": [
            ("IN THE WORD", "The Spirit lifts the veil, reveals Christ, and enables transforming understanding.", ["How do I prepare my heart before reading?", "What passage has the Spirit recently illuminated?", "What response did that illumination require?"]),
            ("IN PRAYER", "The Spirit helps our weakness, aligns prayer with God's will, and draws us into communion.", ["Where does my prayer life depend mostly on effort?", "What burden may the Spirit be inviting me to carry?", "How can I create unhurried prayer space?"]),
            ("IN COMMUNITY", "The Spirit often guides, strengthens, and confirms through members of Christ's body.", ["Who helps me discern God's activity?", "Where have I isolated myself?", "How can I contribute to another believer's growth?"]),
        ],
        "meditation_ref": "John 16:13",
        "meditation_text": "Howbeit when he, the Spirit of truth, is come, he will guide you into all truth.",
        "observe": ["What name is given to the Spirit?", "What ministry does He perform?", "Where do I need guidance into truth?", "What posture supports attentive fellowship?"],
        "story": ["Describe a time the Spirit made Scripture alive to you.", "How have you experienced His help in prayer?", "What fruit has grown through spiritual community?"],
        "practices": ["Begin each study time by welcoming the Spirit's help.", "Set aside twenty minutes for listening prayer.", "Pray one biblical passage back to God.", "Contact a mature believer for prayer and discernment.", "Respond promptly to one tested prompting to serve."],
        "declarations": ["The Holy Spirit dwells in me.", "I welcome His communion.", "He guides me into truth.", "He helps my weakness in prayer.", "I test His leading through Scripture.", "I receive grace through Christ's body.", "Fellowship produces transformation.", "I walk in joyful partnership with the Spirit."],
        "checkpoint": ["I understand the Spirit's role in fellowship.", "I regularly ask for illumination in Scripture.", "I depend on the Spirit in prayer.", "I test and respond to His leading.", "I participate meaningfully in Christian community."],
        "milestone": "The clearest way the Holy Spirit is teaching me to fellowship with God is…",
        "different": "One rhythm of communion I will establish beginning this week is…",
    },
    {
        "num": 9,
        "title": "The Manifestation of the Sons of God",
        "display": "THE MANIFESTATION\nOF THE SONS OF GOD",
        "key_ref": "Romans 8:19",
        "key_text": "For the earnest expectation of the creature waiteth for the manifestation of the sons of God.",
        "supporting": ["Romans 8:18–23", "1 Corinthians 12:7–11", "Ephesians 4:11–13", "Matthew 5:14–16"],
        "objectives": [
            "Understand manifestation as the visible expression of God's character and power through mature sons.",
            "Recognize spiritual gifts as channels for serving real human need.",
            "Connect calling, stewardship, service, and multiplication.",
            "Develop a practical plan for revealing Christ in your sphere of influence.",
        ],
        "prayer": "Father, mature me beyond private spirituality. Let Your character, wisdom, compassion, and power become visible through my service. Show me the needs You have equipped me to address, and make my life fruitful for Your glory. Amen.",
        "focus": ["Expression", "Gifts", "Service", "Multiplication"],
        "pathway": ["Formation", "Alignment", "Empowerment", "Service", "Multiplication"],
        "inventory": [
            "My private formation is producing visible service.",
            "I can name gifts, burdens, and abilities God has entrusted to me.",
            "I look for real needs rather than platforms for recognition.",
            "I depend on the Spirit rather than natural ability alone.",
            "I am helping other believers grow and serve.",
        ],
        "areas": [
            ("CHARACTER REVEALED", "Manifestation begins with God's nature becoming visible through mature conduct and relationships.", ["Where should people see more of Christ in me?", "What pressure exposes unfinished formation?", "What fruit would make God tangible there?"]),
            ("GIFTS RELEASED", "The manifestation of the Spirit is given to profit others and bring divine help into real situations.", ["Which gifts or abilities repeatedly serve others?", "What need stirs holy compassion in me?", "How can I grow in skill, humility, and dependence?"]),
            ("SONS MULTIPLIED", "Mature sons do not merely display gifts; they equip, disciple, and release others.", ["Who am I intentionally strengthening?", "What have I learned that should be passed on?", "What structure could help fruit continue beyond me?"]),
        ],
        "meditation_ref": "Matthew 5:16",
        "meditation_text": "Let your light so shine before men, that they may see your good works, and glorify your Father which is in heaven.",
        "observe": ["What should become visible?", "Who should receive glory?", "What good work is currently before me?", "How can I serve without making myself the center?"],
        "story": ["Where has God already used your life to help another person?", "What gift or burden is becoming clearer?", "What sphere is waiting for a mature expression of Christ through you?"],
        "practices": ["List your gifts, burdens, experiences, and present opportunities.", "Serve one concrete need without seeking recognition.", "Ask three trusted people where they see grace on your life.", "Develop a thirty-day service experiment.", "Begin intentionally discipling or mentoring one person."],
        "declarations": ["Creation will see Christ through surrendered sons.", "God's character is becoming visible in me.", "The Spirit equips me to serve.", "My gifts exist for the profit of others.", "I carry divine solutions into human need.", "I reject platform-centered ministry.", "I will equip and strengthen others.", "My life will bring glory to the Father."],
        "checkpoint": ["I understand biblical manifestation.", "I can name several gifts and burdens entrusted to me.", "I serve real needs with humility.", "I depend on the Spirit in ministry.", "I am beginning to multiply what God has formed in me."],
        "milestone": "The sphere in which God is calling me to reveal Christ more clearly is…",
        "different": "One concrete act of Spirit-led service I will begin now is…",
    },
]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = GOLD, size: str = "8") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, end])


def configure_section(section) -> None:
    section.page_width = Inches(7)
    section.page_height = Inches(10)
    section.top_margin = Inches(0.62)
    section.bottom_margin = Inches(0.62)
    section.left_margin = Inches(0.68)
    section.right_margin = Inches(0.68)
    section.header_distance = Inches(0.25)
    section.footer_distance = Inches(0.3)


def style_run(run, size=10.5, bold=False, italic=False, color=INK_RGB, font="Garamond") -> None:
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color


def add_title(doc, text: str, level=1, color=INK_RGB, align=WD_ALIGN_PARAGRAPH.LEFT, before=0, after=8):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    sizes = {1: 24, 2: 17, 3: 12}
    style_run(r, sizes.get(level, 12), True, False, color, "Georgia")
    return p


def add_body(doc, text: str, size=10.5, bold=False, italic=False, color=INK_RGB, align=None, after=5):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.08
    r = p.add_run(text)
    style_run(r, size, bold, italic, color)
    return p


def add_kicker(doc, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text.upper())
    style_run(r, 8.5, True, False, GOLD_RGB, "Arial")
    r.font.all_caps = True
    return p


def add_lines(doc, count=8, prompt=None):
    if prompt:
        add_body(doc, prompt, 9.5, False, False, INK_RGB, after=3)
    for _ in range(count):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 0.9
        r = p.add_run("________________________________________________________________________________")
        style_run(r, 8, False, False, RGBColor(170, 164, 152), "Arial")


def add_checkbox_list(doc, items, numbered=False):
    for idx, item in enumerate(items, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.05)
        p.paragraph_format.first_line_indent = Inches(-0.02)
        p.paragraph_format.space_after = Pt(4)
        prefix = f"{idx}. " if numbered else "☐  "
        r = p.add_run(prefix + item)
        style_run(r, 10, False, False, INK_RGB)


def boxed_text(doc, heading, text, fill=CREAM, border=GOLD):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Inches(5.45)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_shading(cell, fill)
    set_cell_border(cell, border, "10")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(heading.upper())
    style_run(r, 10, True, False, NAVY_RGB, "Georgia")
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r2 = p2.add_run(text)
    style_run(r2, 9.5, False, True, INK_RGB)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def page_break(doc):
    doc.add_page_break()


def add_docx_header_footer(section):
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("THE DIVINE BLUEPRINT COMPANION")
    style_run(r, 7.5, True, False, GOLD_RGB, "Arial")
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = fp.add_run("THE GLEANING GROUND   •   ")
    style_run(r1, 7.5, True, False, MUTED_RGB, "Arial")
    add_field(fp, "PAGE")


def build_docx() -> None:
    doc = Document()
    for section in doc.sections:
        configure_section(section)
        section.different_first_page_header_footer = True
        add_docx_header_footer(section)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Garamond"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK_RGB
    normal.paragraph_format.space_after = Pt(5)

    # Cover
    cover = doc.add_table(rows=1, cols=1)
    cover.alignment = WD_TABLE_ALIGNMENT.CENTER
    cover.autofit = False
    row = cover.rows[0]
    row.height = Inches(8.62)
    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    cell = cover.cell(0, 0)
    cell.width = Inches(5.64)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_shading(cell, NAVY)
    cp = cell.paragraphs[0]
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cp.add_run("THE DIVINE\nBLUEPRINT")
    style_run(r, 31, True, False, WHITE_RGB, "Georgia")
    cp2 = cell.add_paragraph()
    cp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cp2.add_run("COMPANION JOURNAL")
    style_run(r, 14, True, False, GOLD_RGB, "Arial")
    cp3 = cell.add_paragraph()
    cp3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp3.paragraph_format.space_before = Pt(20)
    r = cp3.add_run("READ  •  REFLECT  •  PRAY  •  PRACTICE  •  BECOME")
    style_run(r, 8.5, True, False, GOLD_RGB, "Arial")
    cp4 = cell.add_paragraph()
    cp4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp4.paragraph_format.space_before = Pt(42)
    r = cp4.add_run("A guided formation journey through\nThe Making, Maturing and Manifestation of God's Sons")
    style_run(r, 11, False, True, CREAM_RGB, "Garamond")
    cp5 = cell.add_paragraph()
    cp5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp5.paragraph_format.space_before = Pt(58)
    r = cp5.add_run("AYO-PAUL IKUJUNI")
    style_run(r, 10, True, False, GOLD_RGB, "Arial")
    cp6 = cell.add_paragraph()
    cp6.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cp6.add_run("THE GLEANING GROUND")
    style_run(r, 8, True, False, WHITE_RGB, "Arial")
    page_break(doc)

    add_kicker(doc, "The Divine Blueprint Companion Journal")
    add_title(doc, "Copyright and Use", 1)
    add_body(doc, "© 2026 Ayo-Paul Ikujuni. All rights reserved.")
    add_body(doc, "This companion journal is designed for use alongside The Divine Blueprint: The Making, Maturing and Manifestation of God's Sons. It may be printed for the purchaser's personal use. It may not be reproduced, resold, uploaded, or distributed without written permission.")
    add_body(doc, "Unless otherwise noted, Scripture quotations in this journal are from the King James Version (KJV).")
    add_body(doc, "A discipleship initiative of The Gleaning Ground.")
    add_body(doc, "Website: divineblueprint.gleaningground.com")
    boxed_text(doc, "Formation, not hurried completion", "This journal is not a test to pass. It is a place to meet God honestly, respond to truth, practice obedience, and record the evidence of His forming work.")
    page_break(doc)

    add_kicker(doc, "How to use this companion")
    add_title(doc, "Read. Reflect. Pray. Practice. Become.", 1)
    for heading, text in [
        ("READ", "Read one chapter of the book slowly. Mark repeated ideas, questions, convictions, and invitations."),
        ("REFLECT", "Use the personal inventory and guided questions to bring the chapter into your actual story."),
        ("PRAY", "Turn insight into conversation with God. Give Him permission to search, correct, heal, and lead."),
        ("PRACTICE", "Choose concrete obedience. Formation deepens when truth is repeatedly practiced."),
        ("BECOME", "Review the spiritual checkpoint, celebrate grace, and name the next area of growth."),
    ]:
        add_kicker(doc, heading)
        add_body(doc, text, 10.5, after=8)
    boxed_text(doc, "Recommended rhythm", "Spend ten days with each chapter. Ninety days creates enough time to read, reflect, practice, revisit, and consolidate the truth without rushing the Spirit's work.")
    page_break(doc)

    add_kicker(doc, "Your ninety-day journey")
    add_title(doc, "A Formation Rhythm for Every Chapter", 1)
    rhythm = [
        "Day 1 — Read the chapter opener, objectives, and preparation prayer.",
        "Day 2 — Read the corresponding chapter in the book with the focus words in mind.",
        "Day 3 — Complete the personal inventory.",
        "Day 4 — Work through reflection area one.",
        "Day 5 — Work through reflection areas two and three.",
        "Day 6 — Complete the Scripture meditation and observation table.",
        "Day 7 — Write your story honestly before God.",
        "Day 8 — Begin the chapter practices and read the declarations aloud.",
        "Day 9 — Use the Kingdom Journal page for listening prayer.",
        "Day 10 — Complete the spiritual checkpoint and blueprint milestone.",
    ]
    add_checkbox_list(doc, rhythm)
    add_title(doc, "Before You Begin", 2)
    add_lines(doc, 5, "What do you hope God will form in you during these ninety days?")
    add_lines(doc, 4, "Who will pray with you or help you remain accountable?")
    page_break(doc)

    add_kicker(doc, "Contents")
    add_title(doc, "The Journey", 1)
    contents = doc.add_table(rows=1, cols=3)
    contents.alignment = WD_TABLE_ALIGNMENT.CENTER
    contents.style = "Table Grid"
    headers = ["CHAPTER", "FORMATION THEME", "DAYS"]
    for i, h in enumerate(headers):
        set_cell_shading(contents.cell(0, i), NAVY)
        p = contents.cell(0, i).paragraphs[0]
        r = p.add_run(h)
        style_run(r, 8.5, True, False, WHITE_RGB, "Arial")
    for ch in CHAPTERS:
        row = contents.add_row().cells
        vals = [f"{ch['num']:02d}", ch["title"], f"{(ch['num']-1)*10+1}–{ch['num']*10}"]
        for i, val in enumerate(vals):
            p = row[i].paragraphs[0]
            r = p.add_run(val)
            style_run(r, 9.5, i == 0, False, INK_RGB)
    doc.add_paragraph()
    add_title(doc, "My Journey Dates", 2)
    for ch in CHAPTERS:
        add_body(doc, f"Chapter {ch['num']:02d}: ____________________ to ____________________", 9.5)
    page_break(doc)

    for ch in CHAPTERS:
        # 1: Opener
        table = doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        left, right = table.rows[0].cells
        left.width = Inches(1.45)
        right.width = Inches(4.15)
        left.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(left, NAVY)
        set_cell_border(left, GOLD, "8")
        set_cell_border(right, GOLD, "8")
        lp = left.paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = lp.add_run("CHAPTER")
        style_run(r, 8.5, True, False, GOLD_RGB, "Arial")
        r = lp.add_run(f"\n{ch['num']:02d}")
        style_run(r, 28, True, False, GOLD_RGB, "Georgia")
        lp2 = left.add_paragraph()
        lp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lp2.paragraph_format.space_before = Pt(155)
        r = lp2.add_run("THE DIVINE\nBLUEPRINT\nCOMPANION")
        style_run(r, 8.2, True, False, GOLD_RGB, "Arial")
        rp = right.paragraphs[0]
        rp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = rp.add_run(ch["display"])
        style_run(r, 25 if len(ch["title"]) < 30 else 20, True, False, NAVY_RGB, "Georgia")
        rp2 = right.add_paragraph()
        rp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rp2.paragraph_format.space_before = Pt(18)
        r = rp2.add_run(f'“{ch["key_text"]}”')
        style_run(r, 11, False, True, INK_RGB)
        rp3 = right.add_paragraph()
        rp3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = rp3.add_run("— " + ch["key_ref"])
        style_run(r, 9, True, False, GOLD_RGB, "Arial")
        rp4 = right.add_paragraph()
        rp4.paragraph_format.space_before = Pt(28)
        r = rp4.add_run("KEY SCRIPTURES")
        style_run(r, 8.5, True, False, GOLD_RGB, "Arial")
        for ref in ch["supporting"]:
            pp = right.add_paragraph()
            pp.paragraph_format.space_after = Pt(2)
            rr = pp.add_run("• " + ref)
            style_run(rr, 9.2, False, False, INK_RGB)
        page_break(doc)

        # 2: Objectives and prayer
        add_kicker(doc, f"Chapter {ch['num']:02d}")
        add_title(doc, "Chapter Objective", 1)
        add_body(doc, "By the end of this chapter you will:", 10, italic=True)
        add_checkbox_list(doc, ch["objectives"])
        boxed_text(doc, "Prepare Your Heart", ch["prayer"])
        add_title(doc, "Read with Purpose", 2)
        add_body(doc, "As you read this chapter, look for these four ideas and underline every place they appear:")
        focus_table = doc.add_table(rows=1, cols=4)
        focus_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, word in enumerate(ch["focus"]):
            set_cell_shading(focus_table.cell(0, i), CREAM)
            set_cell_border(focus_table.cell(0, i), GOLD, "6")
            p = focus_table.cell(0, i).paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(word.upper())
            style_run(r, 8.5, True, False, NAVY_RGB, "Arial")
        page_break(doc)

        # 3: Blueprint map and inventory
        add_kicker(doc, "The Divine Blueprint")
        add_title(doc, "God's Pattern", 1)
        path = doc.add_table(rows=1, cols=len(ch["pathway"]))
        path.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, step in enumerate(ch["pathway"]):
            set_cell_shading(path.cell(0, i), NAVY if i in (0, len(ch["pathway"])-1) else CREAM)
            set_cell_border(path.cell(0, i), GOLD, "6")
            p = path.cell(0, i).paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(step.upper())
            style_run(r, 7.8, True, False, WHITE_RGB if i in (0, len(ch["pathway"])-1) else NAVY_RGB, "Arial")
        add_title(doc, "Personal Inventory", 2, before=12)
        add_body(doc, "Rate each statement from 1 (rarely true) to 5 (consistently true).", 9.5, italic=True)
        inv = doc.add_table(rows=1, cols=6)
        inv.alignment = WD_TABLE_ALIGNMENT.CENTER
        inv.style = "Table Grid"
        head = inv.rows[0].cells
        labels = ["STATEMENT", "1", "2", "3", "4", "5"]
        for i, label in enumerate(labels):
            set_cell_shading(head[i], NAVY)
            p = head[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(label)
            style_run(r, 7.5, True, False, WHITE_RGB, "Arial")
        for statement in ch["inventory"]:
            cells = inv.add_row().cells
            p = cells[0].paragraphs[0]
            r = p.add_run(statement)
            style_run(r, 8.6, False, False, INK_RGB)
            for i in range(1, 6):
                p = cells[i].paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run("○")
                style_run(r, 10, False, False, GOLD_RGB, "Arial")
        add_lines(doc, 4, "What does this inventory reveal about your present stage?")
        page_break(doc)

        # 4: Reflection areas 1 and 2
        add_kicker(doc, "Personal Reflection")
        add_title(doc, "Examine the Work Within", 1)
        for idx, (heading, truth, prompts) in enumerate(ch["areas"][:2], 1):
            add_title(doc, f"{idx}. {heading}", 2, color=NAVY_RGB, before=5, after=3)
            add_body(doc, truth, 9.5, italic=True, color=MUTED_RGB)
            for prompt in prompts:
                add_lines(doc, 3, prompt)
        page_break(doc)

        # 5: Reflection area 3 and chapter synthesis
        add_kicker(doc, "Personal Reflection")
        heading, truth, prompts = ch["areas"][2]
        add_title(doc, f"3. {heading}", 1)
        add_body(doc, truth, 10, italic=True, color=MUTED_RGB)
        for prompt in prompts:
            add_lines(doc, 4, prompt)
        boxed_text(doc, "Chapter Synthesis", "Complete this sentence prayerfully: The central invitation of this chapter for my life is…")
        add_lines(doc, 6)
        page_break(doc)

        # 6: Scripture meditation
        add_kicker(doc, "Scripture Meditation")
        add_title(doc, ch["meditation_ref"], 1)
        boxed_text(doc, "Read Slowly", ch["meditation_text"])
        add_lines(doc, 7, "Write the verse below in your own handwriting.")
        add_lines(doc, 6, "Now rewrite the verse in your own words.")
        add_lines(doc, 5, "What is God emphasizing to you?")
        page_break(doc)

        # 7: Observe, interpret, apply
        add_kicker(doc, "Study the Word")
        add_title(doc, "Observe. Understand. Respond.", 1)
        study = doc.add_table(rows=1, cols=2)
        study.alignment = WD_TABLE_ALIGNMENT.CENTER
        study.style = "Table Grid"
        for i, label in enumerate(["QUESTION", "MY NOTES"]):
            set_cell_shading(study.cell(0, i), NAVY)
            p = study.cell(0, i).paragraphs[0]
            r = p.add_run(label)
            style_run(r, 8.5, True, False, WHITE_RGB, "Arial")
        for question in ch["observe"]:
            cells = study.add_row().cells
            p = cells[0].paragraphs[0]
            r = p.add_run(question)
            style_run(r, 9, False, False, INK_RGB)
            p = cells[1].paragraphs[0]
            for _ in range(4):
                r = p.add_run("________________________________\n")
                style_run(r, 8, False, False, RGBColor(175, 170, 160), "Arial")
        add_lines(doc, 5, "What is one truth to believe, one attitude to change, and one action to take?")
        page_break(doc)

        # 8: My story
        add_kicker(doc, "My Story")
        add_title(doc, "Trace God's Work in Your Life", 1)
        add_body(doc, "Write honestly. You are not composing a testimony for an audience; you are noticing the faithful work of God.", 9.5, italic=True, color=MUTED_RGB)
        for prompt in ch["story"]:
            add_lines(doc, 8, prompt)
        page_break(doc)

        # 9: Practice and declarations
        add_kicker(doc, "Put Truth into Practice")
        add_title(doc, "Practice", 1)
        add_checkbox_list(doc, ch["practices"])
        add_title(doc, "Declarations", 2, before=15)
        add_body(doc, "Read aloud slowly. Do not use these words to avoid honest struggle; use them to bring your mind and speech into agreement with God's truth.", 9, italic=True, color=MUTED_RGB)
        for declaration in ch["declarations"]:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            r = p.add_run("✦  " + declaration)
            style_run(r, 10, False, False, INK_RGB)
        page_break(doc)

        # 10: Kingdom journal and checkpoint
        add_kicker(doc, "Kingdom Journal")
        add_title(doc, "Prompt", 1)
        add_body(doc, "Lord, what are You saying to me after this chapter?", 10.5, italic=True)
        add_lines(doc, 12)
        add_title(doc, "Spiritual Checkpoint", 2, before=10)
        cp = doc.add_table(rows=1, cols=6)
        cp.alignment = WD_TABLE_ALIGNMENT.CENTER
        cp.style = "Table Grid"
        headers = ["STATEMENT", "1", "2", "3", "4", "5"]
        for i, label in enumerate(headers):
            set_cell_shading(cp.cell(0, i), NAVY)
            p = cp.cell(0, i).paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(label)
            style_run(r, 7.3, True, False, WHITE_RGB, "Arial")
        for statement in ch["checkpoint"]:
            cells = cp.add_row().cells
            p = cells[0].paragraphs[0]
            r = p.add_run(statement)
            style_run(r, 8.2, False, False, INK_RGB)
            for i in range(1, 6):
                p = cells[i].paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run("○")
                style_run(r, 9, False, False, GOLD_RGB, "Arial")
        add_lines(doc, 3, ch["milestone"])
        add_lines(doc, 3, ch["different"])
        add_body(doc, "Date: __________________________", 9.5, bold=True)
        page_break(doc)

    add_kicker(doc, "The Journey Continues")
    add_title(doc, "Ninety-Day Review", 1)
    add_lines(doc, 7, "What has God formed in me during this journey?")
    add_lines(doc, 7, "Which chapter produced the deepest change, and why?")
    add_lines(doc, 7, "Where do I still need patient formation?")
    add_lines(doc, 6, "Who can I serve, strengthen, or disciple with what I have learned?")
    page_break(doc)

    add_kicker(doc, "My Rule of Life")
    add_title(doc, "Rhythms That Will Sustain Formation", 1)
    for label in ["SCRIPTURE", "PRAYER", "FELLOWSHIP", "SERVICE", "REST", "ACCOUNTABILITY"]:
        add_title(doc, label, 3, color=GOLD_RGB, before=5, after=2)
        add_lines(doc, 4)
    page_break(doc)

    final = doc.add_table(rows=1, cols=1)
    final.alignment = WD_TABLE_ALIGNMENT.CENTER
    row = final.rows[0]
    row.height = Inches(8.25)
    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    cell = final.cell(0, 0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_shading(cell, NAVY)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("ARISE, SHINE")
    style_run(r, 28, True, False, GOLD_RGB, "Georgia")
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p2.add_run("for thy light is come, and the glory of the LORD is risen upon thee.")
    style_run(r, 13, False, True, WHITE_RGB)
    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p3.add_run("ISAIAH 60:1")
    style_run(r, 9, True, False, GOLD_RGB, "Arial")
    p4 = cell.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p4.paragraph_format.space_before = Pt(70)
    r = p4.add_run("THE DIVINE BLUEPRINT COMPANION\nA discipleship initiative of The Gleaning Ground")
    style_run(r, 8.5, True, False, CREAM_RGB, "Arial")

    core = doc.core_properties
    core.title = "The Divine Blueprint Companion Journal"
    core.subject = "A guided formation journal for The Divine Blueprint"
    core.author = "Ayo-Paul Ikujuni"
    core.keywords = "Christian discipleship, spiritual formation, companion journal, sonship"
    doc.save(DOCX_PATH)


# ---------- PDF helpers ----------

PDF_NAVY = colors.HexColor("#0B2036")
PDF_GOLD = colors.HexColor("#B8862D")
PDF_CREAM = colors.HexColor("#F7F1E5")
PDF_INK = colors.HexColor("#16283A")
PDF_MUTED = colors.HexColor("#6D6A63")
PDF_LINE = colors.HexColor("#C8C0B0")

SERIF = "Times-Roman"
SERIF_BOLD = "Times-Bold"
SERIF_ITALIC = "Times-Italic"
SANS = "Helvetica"
SANS_BOLD = "Helvetica-Bold"


def register_pdf_fonts():
    global SERIF, SERIF_BOLD, SERIF_ITALIC, SANS, SANS_BOLD
    candidates = {
        "TGGSerif": "/usr/share/fonts/truetype/dejavu/DejaVuSerif.ttf",
        "TGGSerifBold": "/usr/share/fonts/truetype/dejavu/DejaVuSerif-Bold.ttf",
        "TGGSerifItalic": "/usr/share/fonts/truetype/dejavu/DejaVuSerif-Italic.ttf",
        "TGGSans": "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "TGGSansBold": "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",
    }
    if all(Path(p).exists() for p in candidates.values()):
        for name, path in candidates.items():
            pdfmetrics.registerFont(TTFont(name, path))
        SERIF, SERIF_BOLD, SERIF_ITALIC = "TGGSerif", "TGGSerifBold", "TGGSerifItalic"
        SANS, SANS_BOLD = "TGGSans", "TGGSansBold"


def wrap_pdf(text, font, size, max_width):
    words = text.split()
    lines = []
    current = ""
    for word in words:
        trial = word if not current else current + " " + word
        if pdfmetrics.stringWidth(trial, font, size) <= max_width:
            current = trial
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def pdf_text(c, text, x, y, width, font=SERIF, size=10, color=PDF_INK, leading=None, align="left"):
    leading = leading or size * 1.35
    lines = []
    for paragraph in str(text).split("\n"):
        lines.extend(wrap_pdf(paragraph, font, size, width) or [""])
    c.setFillColor(color)
    c.setFont(font, size)
    for line in lines:
        if align == "center":
            c.drawCentredString(x + width / 2, y, line)
        elif align == "right":
            c.drawRightString(x + width, y, line)
        else:
            c.drawString(x, y, line)
        y -= leading
    return y


def pdf_kicker(c, text, x=0.62*inch, y=9.25*inch):
    c.setFillColor(PDF_GOLD)
    c.setFont(SANS_BOLD, 7.5)
    c.drawString(x, y, text.upper())
    c.setStrokeColor(PDF_GOLD)
    c.setLineWidth(0.6)
    c.line(x, y-5, PAGE_W-0.62*inch, y-5)
    return y - 24


def pdf_title(c, text, y, size=21, x=0.62*inch, width=PAGE_W-1.24*inch, color=PDF_NAVY, center=False):
    return pdf_text(c, text, x, y, width, SERIF_BOLD, size, color, size*1.12, "center" if center else "left")


def pdf_footer(c, page_num):
    c.setStrokeColor(PDF_LINE)
    c.setLineWidth(0.4)
    c.line(0.62*inch, 0.48*inch, PAGE_W-0.62*inch, 0.48*inch)
    c.setFillColor(PDF_MUTED)
    c.setFont(SANS, 6.8)
    c.drawString(0.62*inch, 0.28*inch, "THE DIVINE BLUEPRINT COMPANION")
    c.drawRightString(PAGE_W-0.62*inch, 0.28*inch, str(page_num))


def pdf_new(c, page_num, kicker=None, dark=False):
    if dark:
        c.setFillColor(PDF_NAVY)
        c.rect(0, 0, PAGE_W, PAGE_H, fill=1, stroke=0)
    else:
        c.setFillColor(PDF_CREAM)
        c.rect(0, 0, PAGE_W, PAGE_H, fill=1, stroke=0)
        pdf_footer(c, page_num)
    return pdf_kicker(c, kicker) if kicker and not dark else 9.15*inch


def pdf_lines(c, y, count, x=0.68*inch, width=PAGE_W-1.36*inch, gap=18):
    c.setStrokeColor(PDF_LINE)
    c.setLineWidth(0.45)
    for _ in range(count):
        c.line(x, y, x+width, y)
        y -= gap
    return y


def pdf_checkbox(c, x, y, label, width, size=9):
    c.setStrokeColor(PDF_GOLD)
    c.setLineWidth(0.7)
    c.rect(x, y-7, 8, 8, fill=0, stroke=1)
    y2 = pdf_text(c, label, x+16, y, width-16, SERIF, size, PDF_INK, size*1.3)
    return min(y-14, y2-2)


def pdf_box(c, x, y_top, width, height, heading, text):
    c.setFillColor(colors.HexColor("#FBF7EF"))
    c.setStrokeColor(PDF_GOLD)
    c.setLineWidth(0.9)
    c.roundRect(x, y_top-height, width, height, 4, fill=1, stroke=1)
    c.setFillColor(PDF_NAVY)
    c.setFont(SANS_BOLD, 8)
    c.drawCentredString(x+width/2, y_top-18, heading.upper())
    pdf_text(c, text, x+18, y_top-36, width-36, SERIF_ITALIC, 8.8, PDF_INK, 12)


def pdf_table_rating(c, y, statements, title="STATEMENT"):
    x = 0.58*inch
    w = PAGE_W - 1.16*inch
    first = w * 0.67
    col = (w-first)/5
    row_h = 0.42*inch
    total_rows = len(statements)+1
    c.setStrokeColor(PDF_LINE)
    c.setLineWidth(0.45)
    c.setFillColor(PDF_NAVY)
    c.rect(x, y-row_h, w, row_h, fill=1, stroke=1)
    headers = [title, "1", "2", "3", "4", "5"]
    positions = [x, x+first, x+first+col, x+first+2*col, x+first+3*col, x+first+4*col]
    widths = [first, col, col, col, col, col]
    c.setFillColor(colors.white)
    c.setFont(SANS_BOLD, 6.5)
    for i, h in enumerate(headers):
        c.drawCentredString(positions[i]+widths[i]/2, y-row_h+10, h)
    y -= row_h
    for statement in statements:
        c.setFillColor(colors.HexColor("#FCF9F2"))
        c.rect(x, y-row_h, w, row_h, fill=1, stroke=1)
        c.setFillColor(PDF_INK)
        c.setFont(SERIF, 6.8)
        lines = wrap_pdf(statement, SERIF, 6.8, first-10)[:3]
        ty = y-10
        for line in lines:
            c.drawString(x+5, ty, line)
            ty -= 8
        for i in range(5):
            cx = x+first+col*i+col/2
            c.setStrokeColor(PDF_GOLD)
            c.circle(cx, y-row_h/2, 3.2, fill=0, stroke=1)
        y -= row_h
    return y


def blueprint_grid(c, x, y, w, h):
    c.saveState()
    c.setStrokeColor(colors.Color(0.72, 0.53, 0.18, alpha=0.23))
    c.setLineWidth(0.25)
    step = 0.25*inch
    xx = x
    while xx <= x+w:
        c.line(xx, y, xx, y+h)
        xx += step
    yy = y
    while yy <= y+h:
        c.line(x, yy, x+w, yy)
        yy += step
    c.circle(x+w/2, y+h/2, min(w,h)*0.22, fill=0, stroke=1)
    c.circle(x+w/2, y+h/2, min(w,h)*0.12, fill=0, stroke=1)
    c.line(x+w/2, y+0.08*h, x+w/2, y+0.92*h)
    c.line(x+0.08*w, y+h/2, x+0.92*w, y+h/2)
    c.restoreState()


def build_pdf() -> None:
    register_pdf_fonts()
    c = canvas.Canvas(str(PDF_PATH), pagesize=(PAGE_W, PAGE_H), pageCompression=1)
    c.setTitle("The Divine Blueprint Companion Journal")
    c.setAuthor("Ayo-Paul Ikujuni")
    c.setSubject("A guided formation journal for The Divine Blueprint")
    c.setKeywords("Christian discipleship, spiritual formation, companion journal, sonship")
    page = 1

    # Cover
    c.setFillColor(PDF_NAVY)
    c.rect(0, 0, PAGE_W, PAGE_H, fill=1, stroke=0)
    blueprint_grid(c, 0.35*inch, 0.35*inch, PAGE_W-0.7*inch, PAGE_H-0.7*inch)
    c.setFillColor(PDF_GOLD)
    c.setFont(SANS_BOLD, 8)
    c.drawCentredString(PAGE_W/2, 8.8*inch, "THE GLEANING GROUND PRESENTS")
    c.setFillColor(colors.white)
    c.setFont(SERIF_BOLD, 30)
    c.drawCentredString(PAGE_W/2, 7.25*inch, "THE DIVINE")
    c.drawCentredString(PAGE_W/2, 6.75*inch, "BLUEPRINT")
    c.setFillColor(PDF_GOLD)
    c.setFont(SANS_BOLD, 14)
    c.drawCentredString(PAGE_W/2, 6.15*inch, "COMPANION JOURNAL")
    c.setStrokeColor(PDF_GOLD)
    c.setLineWidth(1)
    c.line(1.45*inch, 5.73*inch, PAGE_W-1.45*inch, 5.73*inch)
    c.setFillColor(PDF_CREAM)
    c.setFont(SERIF_ITALIC, 10.5)
    c.drawCentredString(PAGE_W/2, 5.35*inch, "Read. Reflect. Pray. Practice. Become.")
    c.setFillColor(PDF_GOLD)
    c.setFont(SANS_BOLD, 8)
    c.drawCentredString(PAGE_W/2, 1.35*inch, "AYO-PAUL IKUJUNI")
    c.setFillColor(colors.white)
    c.setFont(SANS, 7)
    c.drawCentredString(PAGE_W/2, 1.05*inch, "A DISCIPLESHIP INITIATIVE OF THE GLEANING GROUND")
    c.showPage(); page += 1

    y = pdf_new(c, page, "Copyright and Use")
    y = pdf_title(c, "The Divine Blueprint Companion Journal", y)
    y = pdf_text(c, "© 2026 Ayo-Paul Ikujuni. All rights reserved.", 0.68*inch, y-8, PAGE_W-1.36*inch, SERIF, 10)
    y = pdf_text(c, "This companion journal is designed for use alongside The Divine Blueprint: The Making, Maturing and Manifestation of God's Sons. It may be printed for the purchaser's personal use. It may not be reproduced, resold, uploaded, or distributed without written permission.", 0.68*inch, y-12, PAGE_W-1.36*inch, SERIF, 10, PDF_INK, 14)
    y = pdf_text(c, "Unless otherwise noted, Scripture quotations are from the King James Version (KJV).", 0.68*inch, y-10, PAGE_W-1.36*inch, SERIF, 10)
    y = pdf_text(c, "A discipleship initiative of The Gleaning Ground.\ndivineblueprint.gleaningground.com", 0.68*inch, y-20, PAGE_W-1.36*inch, SERIF, 10)
    pdf_box(c, 0.7*inch, y-25, PAGE_W-1.4*inch, 1.45*inch, "Formation, not hurried completion", "This journal is not a test to pass. It is a place to meet God honestly, respond to truth, practice obedience, and record the evidence of His forming work.")
    c.showPage(); page += 1

    y = pdf_new(c, page, "How to Use This Companion")
    y = pdf_title(c, "Read. Reflect. Pray. Practice. Become.", y, 19)
    steps = [
        ("READ", "Read one chapter slowly. Mark repeated ideas, questions, convictions, and invitations."),
        ("REFLECT", "Use the inventory and guided questions to bring truth into your actual story."),
        ("PRAY", "Turn insight into conversation with God. Give Him permission to search and lead."),
        ("PRACTICE", "Choose concrete obedience. Formation deepens when truth is repeatedly practiced."),
        ("BECOME", "Review the checkpoint, celebrate grace, and name the next area of growth."),
    ]
    for heading, text in steps:
        c.setFillColor(PDF_GOLD); c.setFont(SANS_BOLD, 8); c.drawString(0.72*inch, y, heading)
        y = pdf_text(c, text, 1.42*inch, y, PAGE_W-2.14*inch, SERIF, 9.5, PDF_INK, 13) - 7
    pdf_box(c, 0.72*inch, y-10, PAGE_W-1.44*inch, 1.2*inch, "Recommended Rhythm", "Spend ten days with each chapter. Ninety days creates enough time to read, reflect, practice, revisit, and consolidate truth without rushing the Spirit's work.")
    c.showPage(); page += 1

    y = pdf_new(c, page, "Your Ninety-Day Journey")
    y = pdf_title(c, "A Formation Rhythm for Every Chapter", y, 18)
    rhythm = [
        "Day 1 — Opener, objectives, and preparation prayer.",
        "Day 2 — Read the book chapter with the focus words in mind.",
        "Day 3 — Complete the personal inventory.",
        "Day 4 — Work through reflection area one.",
        "Day 5 — Work through reflection areas two and three.",
        "Day 6 — Complete Scripture meditation and study notes.",
        "Day 7 — Write your story honestly before God.",
        "Day 8 — Begin the practices and read declarations aloud.",
        "Day 9 — Use the Kingdom Journal for listening prayer.",
        "Day 10 — Complete the checkpoint and milestone.",
    ]
    for item in rhythm:
        y = pdf_checkbox(c, 0.72*inch, y-2, item, PAGE_W-1.44*inch, 8.7)
    y -= 8
    c.setFillColor(PDF_NAVY); c.setFont(SERIF_BOLD, 13); c.drawString(0.72*inch, y, "Before You Begin")
    y -= 22
    y = pdf_text(c, "What do you hope God will form in you during these ninety days?", 0.72*inch, y, PAGE_W-1.44*inch, SERIF_ITALIC, 9)
    pdf_lines(c, y-5, 6)
    c.showPage(); page += 1

    y = pdf_new(c, page, "Contents")
    y = pdf_title(c, "The Journey", y)
    c.setFont(SANS_BOLD, 7.5); c.setFillColor(PDF_GOLD)
    c.drawString(0.72*inch, y, "CHAPTER")
    c.drawString(1.55*inch, y, "FORMATION THEME")
    c.drawRightString(PAGE_W-0.72*inch, y, "DAYS")
    y -= 14
    c.setStrokeColor(PDF_GOLD); c.line(0.72*inch, y, PAGE_W-0.72*inch, y); y -= 18
    for ch in CHAPTERS:
        c.setFillColor(PDF_GOLD); c.setFont(SERIF_BOLD, 12); c.drawString(0.72*inch, y, f"{ch['num']:02d}")
        c.setFillColor(PDF_INK); c.setFont(SERIF, 9.5); c.drawString(1.55*inch, y, ch["title"])
        c.setFont(SANS, 7.5); c.drawRightString(PAGE_W-0.72*inch, y, f"{(ch['num']-1)*10+1}–{ch['num']*10}")
        y -= 31
    c.showPage(); page += 1

    for ch in CHAPTERS:
        # 1 opener
        c.setFillColor(PDF_CREAM); c.rect(0, 0, PAGE_W, PAGE_H, fill=1, stroke=0)
        c.setFillColor(PDF_NAVY); c.rect(0, 0, 1.5*inch, PAGE_H, fill=1, stroke=0)
        blueprint_grid(c, 0.12*inch, 0.3*inch, 1.26*inch, PAGE_H-0.6*inch)
        c.setFillColor(PDF_GOLD); c.setFont(SANS_BOLD, 7); c.drawCentredString(0.75*inch, 9.2*inch, "CHAPTER")
        c.setFont(SERIF_BOLD, 29); c.drawCentredString(0.75*inch, 8.72*inch, f"{ch['num']:02d}")
        c.setFont(SANS_BOLD, 6.5); c.drawCentredString(0.75*inch, 0.95*inch, "THE DIVINE")
        c.drawCentredString(0.75*inch, 0.79*inch, "BLUEPRINT")
        c.drawCentredString(0.75*inch, 0.63*inch, "COMPANION")
        title_size = 23 if len(ch["title"]) < 25 else 18
        y = pdf_text(c, ch["display"], 1.8*inch, 7.95*inch, PAGE_W-2.15*inch, SERIF_BOLD, title_size, PDF_NAVY, title_size*1.14, "center")
        c.setStrokeColor(PDF_GOLD); c.line(2.35*inch, y-6, PAGE_W-0.85*inch, y-6)
        y = pdf_text(c, f'“{ch["key_text"]}”', 1.9*inch, y-34, PAGE_W-2.3*inch, SERIF_ITALIC, 10, PDF_INK, 14, "center")
        c.setFillColor(PDF_GOLD); c.setFont(SANS_BOLD, 7.2); c.drawCentredString((1.5*inch+PAGE_W)/2, y-5, "— " + ch["key_ref"])
        y -= 47
        c.setFillColor(PDF_GOLD); c.setFont(SANS_BOLD, 7); c.drawString(1.88*inch, y, "KEY SCRIPTURES")
        y -= 16
        for ref in ch["supporting"]:
            c.setFillColor(PDF_INK); c.setFont(SERIF, 8.5); c.drawString(1.92*inch, y, "•  " + ref); y -= 18
        c.showPage(); page += 1

        # 2 objectives
        y = pdf_new(c, page, f"Chapter {ch['num']:02d} • Chapter Objective")
        y = pdf_title(c, "By the end of this chapter you will:", y, 16)
        for item in ch["objectives"]:
            y = pdf_checkbox(c, 0.72*inch, y-4, item, PAGE_W-1.44*inch, 8.7)
        pdf_box(c, 0.7*inch, y-12, PAGE_W-1.4*inch, 1.58*inch, "Prepare Your Heart", ch["prayer"])
        y -= 1.92*inch
        c.setFillColor(PDF_NAVY); c.setFont(SERIF_BOLD, 13); c.drawCentredString(PAGE_W/2, y, "READ WITH PURPOSE")
        y -= 20
        pdf_text(c, "Look for these four ideas and underline every place they appear.", 0.8*inch, y, PAGE_W-1.6*inch, SERIF, 8.5, PDF_INK, 11, "center")
        y -= 42
        box_w = (PAGE_W-1.5*inch)/4
        for i, word in enumerate(ch["focus"]):
            x = 0.75*inch + i*box_w
            c.setStrokeColor(PDF_GOLD); c.setLineWidth(0.7); c.roundRect(x+4, y-25, box_w-8, 34, 3, fill=0, stroke=1)
            c.setFillColor(PDF_NAVY); c.setFont(SANS_BOLD, 6.8); c.drawCentredString(x+box_w/2, y-9, word.upper())
        c.showPage(); page += 1

        # 3 map/inventory
        y = pdf_new(c, page, f"Chapter {ch['num']:02d} • God's Pattern")
        y = pdf_title(c, "The Divine Blueprint", y, 19)
        c.setStrokeColor(PDF_GOLD); c.setLineWidth(1)
        x0 = 0.82*inch; x1 = PAGE_W-0.82*inch
        step = (x1-x0)/(len(ch["pathway"])-1)
        c.line(x0, y-28, x1, y-28)
        for i, step_name in enumerate(ch["pathway"]):
            x = x0 + i*step
            c.setFillColor(PDF_NAVY if i in (0, len(ch["pathway"])-1) else PDF_CREAM)
            c.setStrokeColor(PDF_GOLD); c.circle(x, y-28, 12, fill=1, stroke=1)
            c.setFillColor(PDF_INK); c.setFont(SANS_BOLD, 6.2); c.drawCentredString(x, y-50, step_name.upper())
        y -= 92
        c.setFillColor(PDF_NAVY); c.setFont(SERIF_BOLD, 15); c.drawString(0.65*inch, y, "Personal Inventory")
        y -= 18
        c.setFillColor(PDF_MUTED); c.setFont(SERIF_ITALIC, 8); c.drawString(0.65*inch, y, "Rate each statement from 1 (rarely true) to 5 (consistently true).")
        y -= 13
        y = pdf_table_rating(c, y, ch["inventory"])
        y -= 18
        c.setFillColor(PDF_INK); c.setFont(SERIF_ITALIC, 8.5); c.drawString(0.68*inch, y, "What does this inventory reveal about your present stage?")
        pdf_lines(c, y-14, 5)
        c.showPage(); page += 1

        # 4 reflection 1/2
        y = pdf_new(c, page, f"Chapter {ch['num']:02d} • Personal Reflection")
        y = pdf_title(c, "Examine the Work Within", y, 19)
        for idx, (heading, truth, prompts) in enumerate(ch["areas"][:2], 1):
            c.setFillColor(PDF_GOLD); c.setFont(SANS_BOLD, 8); c.drawString(0.68*inch, y, f"{idx}. {heading}")
            y -= 15
            y = pdf_text(c, truth, 0.68*inch, y, PAGE_W-1.36*inch, SERIF_ITALIC, 8.5, PDF_MUTED, 11) - 4
            for prompt in prompts:
                y = pdf_text(c, prompt, 0.68*inch, y, PAGE_W-1.36*inch, SERIF, 8.2, PDF_INK, 10)
                y = pdf_lines(c, y-2, 2, gap=15) - 5
            y -= 5
        c.showPage(); page += 1

        # 5 reflection 3/synthesis
        y = pdf_new(c, page, f"Chapter {ch['num']:02d} • Personal Reflection")
        heading, truth, prompts = ch["areas"][2]
        y = pdf_title(c, f"3. {heading}", y, 20)
        y = pdf_text(c, truth, 0.68*inch, y-2, PAGE_W-1.36*inch, SERIF_ITALIC, 9, PDF_MUTED, 12) - 8
        for prompt in prompts:
            y = pdf_text(c, prompt, 0.68*inch, y, PAGE_W-1.36*inch, SERIF, 8.7, PDF_INK, 11)
            y = pdf_lines(c, y-3, 4, gap=17) - 9
        pdf_box(c, 0.7*inch, y+3, PAGE_W-1.4*inch, 1.18*inch, "Chapter Synthesis", "Complete this sentence prayerfully: The central invitation of this chapter for my life is…")
        c.showPage(); page += 1

        # 6 meditation
        y = pdf_new(c, page, f"Chapter {ch['num']:02d} • Scripture Meditation")
        y = pdf_title(c, ch["meditation_ref"], y, 20)
        pdf_box(c, 0.7*inch, y-4, PAGE_W-1.4*inch, 1.1*inch, "Read Slowly", ch["meditation_text"])
        y -= 1.42*inch
        y = pdf_text(c, "Write the verse below in your own handwriting.", 0.68*inch, y, PAGE_W-1.36*inch, SERIF_ITALIC, 8.5)
        y = pdf_lines(c, y-5, 7, gap=17) - 8
        y = pdf_text(c, "Now rewrite the verse in your own words.", 0.68*inch, y, PAGE_W-1.36*inch, SERIF_ITALIC, 8.5)
        y = pdf_lines(c, y-5, 6, gap=17) - 8
        y = pdf_text(c, "What is God emphasizing to you?", 0.68*inch, y, PAGE_W-1.36*inch, SERIF_ITALIC, 8.5)
        pdf_lines(c, y-5, 5, gap=17)
        c.showPage(); page += 1

        # 7 study table
        y = pdf_new(c, page, f"Chapter {ch['num']:02d} • Study the Word")
        y = pdf_title(c, "Observe. Understand. Respond.", y, 18)
        x = 0.58*inch; w = PAGE_W-1.16*inch; left_w = 2.18*inch; right_w = w-left_w; row_h = 1.2*inch
        c.setFillColor(PDF_NAVY); c.setStrokeColor(PDF_LINE); c.rect(x, y-0.35*inch, w, 0.35*inch, fill=1, stroke=1)
        c.setFillColor(colors.white); c.setFont(SANS_BOLD, 7); c.drawString(x+7, y-0.22*inch, "QUESTION"); c.drawString(x+left_w+7, y-0.22*inch, "MY NOTES")
        y -= 0.35*inch
        for q in ch["observe"]:
            c.setFillColor(colors.HexColor("#FCF9F2")); c.setStrokeColor(PDF_LINE); c.rect(x, y-row_h, w, row_h, fill=1, stroke=1)
            c.line(x+left_w, y, x+left_w, y-row_h)
            pdf_text(c, q, x+7, y-14, left_w-14, SERIF, 8, PDF_INK, 10)
            pdf_lines(c, y-20, 4, x+left_w+8, right_w-16, 15)
            y -= row_h
        y -= 13
        y = pdf_text(c, "One truth to believe, one attitude to change, and one action to take:", 0.68*inch, y, PAGE_W-1.36*inch, SERIF_ITALIC, 8.5)
        pdf_lines(c, y-4, 5, gap=16)
        c.showPage(); page += 1

        # 8 story
        y = pdf_new(c, page, f"Chapter {ch['num']:02d} • My Story")
        y = pdf_title(c, "Trace God's Work in Your Life", y, 18)
        y = pdf_text(c, "Write honestly. You are not composing a testimony for an audience; you are noticing the faithful work of God.", 0.68*inch, y, PAGE_W-1.36*inch, SERIF_ITALIC, 8.3, PDF_MUTED, 11) - 10
        for prompt in ch["story"]:
            y = pdf_text(c, prompt, 0.68*inch, y, PAGE_W-1.36*inch, SERIF, 9, PDF_INK, 12)
            y = pdf_lines(c, y-3, 8, gap=16) - 10
        c.showPage(); page += 1

        # 9 practices/declarations
        y = pdf_new(c, page, f"Chapter {ch['num']:02d} • Put Truth into Practice")
        y = pdf_title(c, "Practice", y, 19)
        for item in ch["practices"]:
            y = pdf_checkbox(c, 0.72*inch, y-2, item, PAGE_W-1.44*inch, 8.6)
        y -= 10
        c.setFillColor(PDF_NAVY); c.setFont(SERIF_BOLD, 15); c.drawString(0.72*inch, y, "Declarations")
        y -= 22
        y = pdf_text(c, "Read aloud slowly. Use these words to bring your mind and speech into agreement with God's truth.", 0.72*inch, y, PAGE_W-1.44*inch, SERIF_ITALIC, 8, PDF_MUTED, 10) - 5
        for declaration in ch["declarations"]:
            c.setFillColor(PDF_GOLD); c.setFont(SANS_BOLD, 7); c.drawString(0.75*inch, y, "✦")
            y = pdf_text(c, declaration, 0.95*inch, y, PAGE_W-1.7*inch, SERIF, 8.7, PDF_INK, 11) - 2
        c.showPage(); page += 1

        # 10 journal/checkpoint
        y = pdf_new(c, page, f"Chapter {ch['num']:02d} • Kingdom Journal")
        y = pdf_title(c, "Prompt", y, 18)
        y = pdf_text(c, "Lord, what are You saying to me after this chapter?", 0.68*inch, y, PAGE_W-1.36*inch, SERIF_ITALIC, 9.5)
        y = pdf_lines(c, y-5, 10, gap=16) - 12
        c.setFillColor(PDF_NAVY); c.setFont(SERIF_BOLD, 14); c.drawString(0.62*inch, y, "Spiritual Checkpoint")
        y -= 16
        y = pdf_table_rating(c, y, ch["checkpoint"])
        y -= 13
        y = pdf_text(c, ch["milestone"], 0.68*inch, y, PAGE_W-1.36*inch, SERIF_ITALIC, 8)
        y = pdf_lines(c, y-2, 2, gap=14) - 6
        y = pdf_text(c, ch["different"], 0.68*inch, y, PAGE_W-1.36*inch, SERIF_ITALIC, 8)
        y = pdf_lines(c, y-2, 2, gap=14)
        c.setFillColor(PDF_MUTED); c.setFont(SANS, 7); c.drawString(0.68*inch, 0.63*inch, "Date: __________________________")
        c.showPage(); page += 1

    y = pdf_new(c, page, "The Journey Continues")
    y = pdf_title(c, "Ninety-Day Review", y, 22)
    for prompt, count in [
        ("What has God formed in me during this journey?", 7),
        ("Which chapter produced the deepest change, and why?", 7),
        ("Where do I still need patient formation?", 7),
        ("Who can I strengthen or disciple with what I have learned?", 6),
    ]:
        y = pdf_text(c, prompt, 0.68*inch, y, PAGE_W-1.36*inch, SERIF_ITALIC, 8.5)
        y = pdf_lines(c, y-3, count, gap=15) - 8
    c.showPage(); page += 1

    y = pdf_new(c, page, "My Rule of Life")
    y = pdf_title(c, "Rhythms That Will Sustain Formation", y, 20)
    for label in ["SCRIPTURE", "PRAYER", "FELLOWSHIP", "SERVICE", "REST", "ACCOUNTABILITY"]:
        c.setFillColor(PDF_GOLD); c.setFont(SANS_BOLD, 7.2); c.drawString(0.68*inch, y, label)
        y = pdf_lines(c, y-7, 4, gap=14) - 8
    c.showPage(); page += 1

    c.setFillColor(PDF_NAVY); c.rect(0, 0, PAGE_W, PAGE_H, fill=1, stroke=0)
    blueprint_grid(c, 0.35*inch, 0.35*inch, PAGE_W-0.7*inch, PAGE_H-0.7*inch)
    c.setFillColor(PDF_GOLD); c.setFont(SERIF_BOLD, 28); c.drawCentredString(PAGE_W/2, 6.5*inch, "ARISE, SHINE")
    c.setFillColor(colors.white); c.setFont(SERIF_ITALIC, 11)
    c.drawCentredString(PAGE_W/2, 5.88*inch, "for thy light is come,")
    c.drawCentredString(PAGE_W/2, 5.63*inch, "and the glory of the LORD is risen upon thee.")
    c.setFillColor(PDF_GOLD); c.setFont(SANS_BOLD, 8); c.drawCentredString(PAGE_W/2, 5.2*inch, "ISAIAH 60:1")
    c.setFillColor(PDF_CREAM); c.setFont(SANS_BOLD, 7.5); c.drawCentredString(PAGE_W/2, 1.15*inch, "THE DIVINE BLUEPRINT COMPANION")
    c.setFont(SANS, 6.5); c.drawCentredString(PAGE_W/2, 0.9*inch, "A DISCIPLESHIP INITIATIVE OF THE GLEANING GROUND")
    c.showPage()
    c.save()


def write_notes() -> None:
    NOTES_PATH.write_text(
        """THE DIVINE BLUEPRINT COMPANION JOURNAL — PRODUCTION NOTES\n\n"
        "FILES\n"
        "1. Editable DOCX: The-Divine-Blueprint-Companion-Journal-Editable.docx\n"
        "2. Print-ready PDF: The-Divine-Blueprint-Companion-Journal-Print-Ready.pdf\n\n"
        "SPECIFICATIONS\n"
        "Trim size: 7 × 10 inches\n"
        "Interior: approximately 97 pages, including front matter, 10 guided pages per chapter, review pages, and closing page\n"
        "Color system: deep navy, muted gold, warm cream\n"
        "PDF fonts: embedded DejaVu fonts when available on the build system; otherwise standard PDF fonts\n"
        "Scripture: King James Version unless otherwise noted\n\n"
        "PRINT GUIDANCE\n"
        "• Suitable for premium workbook or journal printing.\n"
        "• Recommended interior paper: 80–100 gsm cream or natural uncoated stock.\n"
        "• Recommended binding: lay-flat sewn, wire-o, or durable perfect binding.\n"
        "• The current PDF is a finished 7 × 10 interior file. Confirm the printer's bleed, gutter, and blank-page requirements before mass production.\n"
        "• The editable DOCX may reflow slightly when opened with fonts or applications different from Microsoft Word.\n\n"
        "CONTENT DESIGN\n"
        "Each chapter includes: opener, objectives, preparation prayer, focus words, formation pathway, personal inventory, guided reflections, Scripture meditation, observation notes, personal story, practices, declarations, Kingdom Journal, spiritual checkpoint, and milestone.\n"
        """,
        encoding="utf-8",
    )


if __name__ == "__main__":
    build_docx()
    build_pdf()
    write_notes()
    print(f"Created: {DOCX_PATH}")
    print(f"Created: {PDF_PATH}")
    print(f"Created: {NOTES_PATH}")
